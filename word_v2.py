#This python script edits the cover page content
#the script edits line 1 where Lapin AMK -logo is,
#curriculum text and year range at lines 6-7
#deleting confirmation information at lines 11-19
import os
from os.path import join 
from docx import Document
from docx.shared import Pt
import re
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE

#Searches text in word document to be replaced while maintaining style
def replace_string(filename):
    doc = filename
    for p in document.paragraphs:
        if '2020-2024' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '2020-2024' in inline[i].text:
                    #use string variables for text.replace(old text, new text)
                    file_end_year_int = file_start_year_int + 4
                    file_end_year = str(file_end_year_int)
                    text = inline[i].text.replace('2020-2024', file_start_year + '-' + file_end_year)
                    inline[i].text = text
            print (p.text)
        if '2020-2020' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '2020-2020' in inline[i].text:
                    #use string variables for text.replace(old text, new text)
                    text = inline[i].text.replace('2020-2020', file_start_year + '-' + file_start_year)
                    inline[i].text = text
            print (p.text)
        if '2020-2021' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '2020-2021' in inline[i].text:
                    #use string variables for text.replace(old text, new text)
                    file_end_year_int = file_start_year_int + 1
                    file_end_year = str(file_end_year_int)
                    text = inline[i].text.replace('2020-2021', file_start_year + '-' + file_end_year)
                    inline[i].text = text
            print (p.text)
        if '2020-2022' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '2020-2022' in inline[i].text:
                    #use string variables for text.replace(old text, new text)
                    file_end_year_int = file_start_year_int + 2
                    file_end_year = str(file_end_year_int)
                    text = inline[i].text.replace('2020-2022', file_start_year + '-' + file_end_year)
                    inline[i].text = text
            print (p.text)
        if '2020-2023' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if '2020-2023' in inline[i].text:
                    #use string variables for text.replace(old text, new text)
                    file_end_year_int = file_start_year_int + 3
                    file_end_year = str(file_end_year_int)
                    text = inline[i].text.replace('2020-2023', file_start_year + '-' + file_end_year)
                    inline[i].text = text
            print (p.text)
        

    return 1


def insert_paragraph_before(self, text=None, style=None):
    """
    Return a newly created paragraph, inserted directly before this
    paragraph. If *text* is supplied, the new paragraph contains that
    text in a single run. If *style* is provided, that style is assigned
    to the new paragraph.
    """
    paragraph = self._insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        paragraph.add_run(text).bold = True
    if style is not None:
        paragraph.style = style
    return paragraph

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

#deletes paragraphs from line 11 to 19 from the given document
#deletes curriculum section accordingly to ECTS 
def delete_docx_prefix_description(document):
    if file_start_year_int >= 2014:
        if ects_float < 30:
            for x in range(8,18):
                delete_paragraph(all_paragraphs[x])
        else:
            for x in range(10,18):
                delete_paragraph(all_paragraphs[x])  
    else:
        if ects_float < 30:
            for x in range(8,18):
                delete_paragraph(all_paragraphs[x])
                
            delete_paragraph(all_paragraphs[0])
        else:
            for x in range(10,18):
                delete_paragraph(all_paragraphs[x])
                
            delete_paragraph(all_paragraphs[0])
        


#This is an example of the order of workflow if document is pre 2014 
#if the document is from after 2014 insert_paragraph_before function is not used in the workflow 
# replace_string('C:/Users/jyrit/Documents/Word_mass_update_proj/2007/OPS_2007-2008_501R07_2020-05-19.docx')
# insert_paragraph_before(all_paragraphs[1], "Rovaniemen ammattikorkeakoulu", 'Otsikko')
# delete_docx_prefix_description('C:/Users/jyrit/Documents/Word_mass_update_proj/2007/OPS_2007-2008_501R07_2020-05-19.docx')
# document.save('C:/Users/jyrit/Documents/Word_mass_update_proj/2007/OPS_2007-2008_501R07_2020-05-19.docx')


for root, dirs, files in os.walk('.'):
    for filepath in files:
        if '.docx' in filepath:
            filepath = os.path.join(root, filepath)
            #starting to modify document.
            filename = os.path.basename(filepath)
            document = Document(filepath)
            all_paragraphs = document.paragraphs
            #this needs to be replace by the actual filename string 
            filename_to_split = filename
            filename_splitted = re.split('[-._]', filename_to_split)
            #fetch the document year from splitted filename string
            file_start_year = filename_splitted[1]
            file_start_year_int = int(file_start_year)
            #Document ECTS search from paragraph[6]
            #if ECTS is less than 30 remove curriculum section from cover page
            ects = document.paragraphs[6]
            for run in ects.runs:
                ects_splitting = run.text.split()
                ects_string = ects_splitting[-2]
            ects_float = float(ects_string)
            #Adding a custom style to document to place Rovaniemen ammattikorkeakoulu text correctly.
            style = document.styles.add_style('Otsikko', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Arial'
            font.size = Pt(18)
            #Checking if document is before 2014 or after and choosing workflow accordingly
            if file_start_year_int >= 2014:
                replace_string(filepath)
                delete_docx_prefix_description(filepath)
                document.save(filepath)
            else:
                replace_string(filepath)
                insert_paragraph_before(all_paragraphs[1], "Rovaniemen ammattikorkeakoulu", 'Otsikko')
                delete_docx_prefix_description(filepath)
                document.save(filepath)
            